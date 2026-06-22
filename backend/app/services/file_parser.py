"""Extract plain text from uploaded CV files (PDF, DOCX, TXT)."""
import io


class UnsupportedFileError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace")
    raise UnsupportedFileError(f"Unsupported file type: {filename}. Use PDF, DOCX or TXT.")


def _from_pdf(data: bytes) -> str:
    import pdfplumber

    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    text = "\n".join(pages).strip()
    if not text:
        raise UnsupportedFileError("No text could be extracted from this PDF (it may be a scanned image).")
    return text


def _from_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(p for p in parts if p.strip()).strip()
    if not text:
        raise UnsupportedFileError("No text could be extracted from this DOCX file.")
    return text
